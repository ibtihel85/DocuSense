"""
backend/ocr/processor.py
─────────────────────────
Document OCR and text extraction pipeline.
Handles: PDF (text + scanned), DOCX, images.
Uses Tesseract for OCR, pypdf for native PDF text.
"""
from __future__ import annotations
import io
from pathlib import Path
from typing import NamedTuple
import pytesseract
from PIL import Image
from backend.core.config import get_settings
from backend.ocr.cleaner import clean_ocr_text, clean_pdf_text
from backend.utils.logger import get_logger

logger = get_logger(__name__)
settings = get_settings()

# Set Tesseract binary path from config
if settings.tesseract_cmd != "tesseract":
    pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd


class ExtractedPage(NamedTuple):
    page_number: int
    text: str
    method: str  # "native_pdf" | "ocr" | "docx"
    confidence: float  # 0.0-1.0 OCR confidence, 1.0 for native


class DocumentExtractor:
    """
    Unified document text extractor.
    Automatically chooses between native PDF parsing and OCR.
    """

    MIN_NATIVE_CHARS = 50  # If native PDF yields fewer chars → use OCR

    def extract(self, file_path: str | Path) -> list[ExtractedPage]:
        """
        Extract text from a document, returning one ExtractedPage per page.
        Auto-detects file type and chooses best extraction method.
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        logger.info("Extracting text from document", file=path.name, type=suffix)

        if suffix == ".pdf":
            return self._extract_pdf(path)
        elif suffix in (".docx", ".doc"):
            return self._extract_docx(path)
        elif suffix in (".png", ".jpg", ".jpeg", ".tiff", ".bmp"):
            return self._extract_image(path)
        elif suffix == ".txt":
            return self._extract_txt(path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    def _extract_pdf(self, path: Path) -> list[ExtractedPage]:
        """Extract PDF: try native text first, fall back to OCR per page."""
        try:
            import pypdf
        except ImportError:
            raise ImportError("pypdf required: pip install pypdf")

        pages: list[ExtractedPage] = []

        with open(path, "rb") as f:
            reader = pypdf.PdfReader(f)
            for page_num, page in enumerate(reader.pages, start=1):
                native_text = page.extract_text() or ""
                native_text = clean_pdf_text(native_text)

                # Decide: native text good enough?
                if len(native_text.strip()) >= self.MIN_NATIVE_CHARS:
                    pages.append(ExtractedPage(
                        page_number=page_num,
                        text=native_text,
                        method="native_pdf",
                        confidence=1.0,
                    ))
                else:
                    # Fall back to OCR for this page
                    logger.debug("Native PDF text sparse, using OCR", page=page_num)
                    ocr_result = self._ocr_pdf_page(path, page_num)
                    pages.append(ocr_result)

        logger.info("PDF extraction complete", pages=len(pages), file=path.name)
        return pages

    def _ocr_pdf_page(self, pdf_path: Path, page_num: int) -> ExtractedPage:
        """Render a PDF page to image and OCR it."""
        try:
            from pdf2image import convert_from_path
            images = convert_from_path(
                str(pdf_path),
                dpi=settings.ocr_dpi,
                first_page=page_num,
                last_page=page_num,
            )
            if not images:
                return ExtractedPage(page_num, "", "ocr", 0.0)

            img = images[0]
            # Preprocess: convert to grayscale for better OCR
            img = img.convert("L")
            text, confidence = self._run_tesseract(img)
            return ExtractedPage(page_num, clean_ocr_text(text), "ocr", confidence)
        except Exception as e:
            logger.warning("OCR failed for PDF page", page=page_num, error=str(e))
            return ExtractedPage(page_num, "", "ocr", 0.0)

    def _extract_docx(self, path: Path) -> list[ExtractedPage]:
        """Extract text from DOCX, treating the whole file as one 'page'."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx required: pip install python-docx")

        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        full_text = "\n\n".join(paragraphs)
        full_text = clean_pdf_text(full_text)

        return [ExtractedPage(page_number=1, text=full_text, method="docx", confidence=1.0)]

    def _extract_image(self, path: Path) -> list[ExtractedPage]:
        """OCR a standalone image file."""
        img = Image.open(path).convert("L")
        text, confidence = self._run_tesseract(img)
        return [ExtractedPage(page_number=1, text=clean_ocr_text(text), method="ocr", confidence=confidence)]

    def _extract_txt(self, path: Path) -> list[ExtractedPage]:
        """Read plain text file."""
        text = path.read_text(encoding="utf-8", errors="replace")
        return [ExtractedPage(page_number=1, text=text.strip(), method="native_pdf", confidence=1.0)]

    def _run_tesseract(self, img: Image.Image) -> tuple[str, float]:
        """Run Tesseract OCR and return (text, confidence_score)."""
        try:
            # Get text with confidence data
            data = pytesseract.image_to_data(
                img,
                lang=settings.tesseract_lang,
                output_type=pytesseract.Output.DICT,
            )
            # Calculate mean confidence (ignore -1 = no text)
            confidences = [int(c) for c in data["conf"] if str(c) != "-1" and int(c) > 0]
            avg_conf = sum(confidences) / len(confidences) / 100.0 if confidences else 0.0

            text = pytesseract.image_to_string(img, lang=settings.tesseract_lang)
            return text, avg_conf
        except Exception as e:
            logger.warning("Tesseract OCR error", error=str(e))
            return "", 0.0


def get_extractor() -> DocumentExtractor:
    return DocumentExtractor()
